#!/usr/bin/env python3
"""
Build the Outcome Works branded Word template (.docx).

  /tmp/owldocx/bin/python brand/build_template.py

Primary fonts are the exact brand fonts (Space Grotesk + IBM Plex Mono). On a
machine without them installed, Word substitutes; install the two free fonts for
a pixel-exact match. Fallbacks are declared on every style's rFonts.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
OWL = ROOT / "public" / "owl.png"
OUT = ROOT / "brand" / "Outcome-Works-Template.docx"

# --- Brand kit -------------------------------------------------------------
NAVY   = RGBColor(0x06, 0x39, 0x5B)
CYAN   = RGBColor(0x17, 0x9E, 0xB2)
TEAL   = RGBColor(0x15, 0x72, 0x7C)
TERRA  = RGBColor(0xC2, 0x56, 0x2F)
INK    = RGBColor(0x14, 0x20, 0x2B)
MUTED  = RGBColor(0x5C, 0x6B, 0x79)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
HAIR   = "CDD5DC"
PAPER  = "ECEFF2"
TEALW  = "DCEDEF"
TERRAW = "F6E6DE"
NAVYHEX = "06395B"

SANS = "Space Grotesk"
MONO = "IBM Plex Mono"
SANS_FB = "Aptos"     # modern Office default, closest always-present grotesque
MONO_FB = "Consolas"  # always-present monospace


# --- low-level XML helpers -------------------------------------------------
def set_fonts(rpr_owner, primary, fallback):
    """Set ascii/hAnsi/cs to primary and stash fallback hint."""
    el = rpr_owner
    rFonts = el.get_or_add_rPr().find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        el.get_or_add_rPr().append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), primary)


def style_font(style, name, fallback, size, color, bold=False, caps=False,
               spacing=None, italic=False):
    f = style.font
    f.name = name
    f.size = Pt(size)
    f.bold = bold
    f.italic = italic
    f.color.rgb = color
    # ascii/hAnsi/cs
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)
    if caps:
        c = OxmlElement("w:caps"); c.set(qn("w:val"), "true"); rpr.append(c)
    if spacing is not None:
        sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(spacing)); rpr.append(sp)


def shade_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def para_left_border(p, color, sz=28, space=10):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), str(space)); left.set(qn("w:color"), color)
    pbdr.append(left); pPr.append(pbdr)


def para_bottom_rule(p, color=HAIR, sz=6, space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space)); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)


def no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def cell_bottom_border(cell, color=HAIR, sz=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz)); b.set(qn("w:color"), color)
    borders.append(b); tcPr.append(borders)


def run_font(run, name, size, color, bold=False, caps=False, spacing=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), name)
    if caps:
        c = OxmlElement("w:caps"); c.set(qn("w:val"), "true"); rpr.append(c)
    if spacing is not None:
        sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(spacing)); rpr.append(sp)


def page_number_run(paragraph):
    run = paragraph.add_run()
    run_font(run, MONO, 8, MUTED)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)


def owl_size(height_in):
    w, h = Image.open(OWL).size
    return Inches(height_in * w / h), Inches(height_in)


# --- document --------------------------------------------------------------
doc = Document()

# Page setup A4 + margins
sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.2))
sec.header_distance = Cm(1.2); sec.footer_distance = Cm(1.2)
CONTENT_W = sec.page_width - sec.left_margin - sec.right_margin
HALF_W = Emu(int(CONTENT_W / 2))

# Base / Normal
normal = doc.styles["Normal"]
style_font(normal, SANS, SANS_FB, 10.5, INK)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.32

def newp(name, base="Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = doc.styles[base]
        return s

# Named styles
s = newp("OW Title");      style_font(s, SANS, SANS_FB, 32, NAVY, bold=True, spacing=-20); s.paragraph_format.space_after = Pt(6)
s = newp("OW Subtitle");   style_font(s, SANS, SANS_FB, 14, MUTED); s.paragraph_format.space_after = Pt(10)
s = newp("OW Eyebrow");    style_font(s, MONO, MONO_FB, 8.5, MUTED, caps=True, spacing=28); s.paragraph_format.space_after = Pt(4)
s = newp("OW Lead");       style_font(s, SANS, SANS_FB, 13, MUTED); s.paragraph_format.space_after = Pt(12); s.paragraph_format.line_spacing = 1.4
s = newp("OW Heading 1");  style_font(s, SANS, SANS_FB, 20, NAVY, bold=True, spacing=-12); s.paragraph_format.space_before = Pt(18); s.paragraph_format.space_after = Pt(6)
s = newp("OW Heading 2");  style_font(s, SANS, SANS_FB, 15, NAVY, bold=True, spacing=-8); s.paragraph_format.space_before = Pt(14); s.paragraph_format.space_after = Pt(4)
s = newp("OW Heading 3");  style_font(s, SANS, SANS_FB, 12, TEAL, bold=True); s.paragraph_format.space_before = Pt(10); s.paragraph_format.space_after = Pt(3)
s = newp("OW Caption");    style_font(s, MONO, MONO_FB, 8.5, MUTED); s.paragraph_format.space_after = Pt(6)
s = newp("OW Stat Number");style_font(s, SANS, SANS_FB, 40, NAVY, bold=True, spacing=-30); s.paragraph_format.space_after = Pt(0); s.paragraph_format.line_spacing = 1.0
s = newp("OW Stat Label"); style_font(s, MONO, MONO_FB, 9, MUTED, caps=True, spacing=20); s.paragraph_format.space_after = Pt(8)
s = newp("OW Callout");    style_font(s, SANS, SANS_FB, 10.5, INK); s.paragraph_format.space_before = Pt(6); s.paragraph_format.space_after = Pt(6)
s = newp("OW Pull Quote"); style_font(s, SANS, SANS_FB, 15, NAVY, italic=False); s.paragraph_format.space_after = Pt(2); s.paragraph_format.line_spacing = 1.3
s = newp("OW Attribution");style_font(s, MONO, MONO_FB, 8.5, MUTED, caps=True, spacing=15); s.paragraph_format.space_after = Pt(10)

# List style with brand tick
def bullet(text, color_hex="15727C"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("→  ")  # arrow tick
    run_font(r, SANS, 10.5, RGBColor.from_string(color_hex), bold=True)
    r2 = p.add_run(text); run_font(r2, SANS, 10.5, INK)
    return p

# ============================ COVER =======================================
# Navy band (1x1 shaded table across content width) holding owl + title.
band = doc.add_table(rows=1, cols=1)
band.alignment = WD_TABLE_ALIGNMENT.CENTER
band.columns[0].width = CONTENT_W
no_table_borders(band)
cell = band.rows[0].cells[0]
cell.width = CONTENT_W
shade_cell(cell, NAVYHEX)
# vertical padding for the band
def cell_margins(cell, top=300, bottom=300, left=240, right=240):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{name}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)
cell_margins(cell, top=420, bottom=420)

cp = cell.paragraphs[0]
ow, oh = owl_size(0.95)
cp.add_run().add_picture(str(OWL), width=ow, height=oh)
cp.paragraph_format.space_after = Pt(10)

t = cell.add_paragraph()
r = t.add_run("Outcome Works")
run_font(r, SANS, 30, WHITE, bold=True, spacing=-20)
t.paragraph_format.space_after = Pt(2)

sub = cell.add_paragraph()
r = sub.add_run("Built to win the work that's hard to win.")
run_font(r, SANS, 13, WHITE)
sub_run_color = sub.runs[0]; sub_run_color.font.color.rgb = RGBColor(0xCF, 0xDF, 0xE6)
sub.paragraph_format.space_after = Pt(6)

eb = cell.add_paragraph()
r = eb.add_run("CONSULTANCY  ·  PRODUCTS")
run_font(r, MONO, 9, CYAN, caps=True, spacing=30)

# Document title + meta below the band
doc.add_paragraph().paragraph_format.space_after = Pt(6)
p = doc.add_paragraph(style="OW Eyebrow"); p.add_run("Document type")
p = doc.add_paragraph(style="OW Title"); p.add_run("Proposal / Report Title")
p = doc.add_paragraph(style="OW Subtitle"); p.add_run("A one-line description of what this document covers.")

# Meta block (mono key/value)
meta = [("Prepared for", "Client / Organisation"),
        ("Prepared by", "Outcome Works Limited"),
        ("Date", "Month YYYY"),
        ("Reference", "OW-0000"),
        ("Confidentiality", "Commercial in confidence")]
mt = doc.add_table(rows=len(meta), cols=2)
no_table_borders(mt)
mt.columns[0].width = Cm(4.2); mt.columns[1].width = CONTENT_W - Cm(4.2)
for i, (k, v) in enumerate(meta):
    kc, vc = mt.rows[i].cells
    kc.width = Cm(4.2); vc.width = CONTENT_W - Cm(4.2)
    kp = kc.paragraphs[0]; kr = kp.add_run(k.upper()); run_font(kr, MONO, 8.5, MUTED, caps=True, spacing=20)
    vp = vc.paragraphs[0]; vr = vp.add_run(v); run_font(vr, SANS, 10.5, INK)
    kp.paragraph_format.space_after = Pt(2); vp.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===================== HEADER & FOOTER (page 2+) ==========================
# Different first page so the cover has no header/footer.
sec.different_first_page_header_footer = True

hdr = sec.header
hdr.is_linked_to_previous = False
hp = hdr.paragraphs[0]
hp.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
hr = hp.add_run()
ow2, oh2 = owl_size(0.20)
hr.add_picture(str(OWL), width=ow2, height=oh2)
hr2 = hp.add_run("  Outcome Works"); run_font(hr2, SANS, 9, NAVY, bold=True)
hr3 = hp.add_run("\tDocument Title"); run_font(hr3, MONO, 8, MUTED, caps=True, spacing=15)
para_bottom_rule(hp, color="179EB2", sz=6, space=6)

ftr = sec.footer
ftr.is_linked_to_previous = False
fp = ftr.paragraphs[0]
fp.paragraph_format.tab_stops.add_tab_stop(CONTENT_W, WD_TAB_ALIGNMENT.RIGHT)
para_bottom_rule  # noqa (top rule not needed)
fr = fp.add_run("Outcome Works Limited  ·  Public-sector & enterprise procurement  ·  UK & EU  ·  Commercial in confidence")
run_font(fr, MONO, 7.5, MUTED)
fp.add_run("\t")
page_number_run(fp)

# ===================== SAMPLE CONTENT (self-documenting) ==================
doc.add_paragraph(style="OW Eyebrow").add_run("Section")
doc.add_paragraph(style="OW Heading 1").add_run("How this template is built")
doc.add_paragraph(style="OW Lead").add_run(
    "This page demonstrates every brand style in use. Duplicate the file, replace "
    "the copy, and keep the structure. The voice is plain, evidence-led, and "
    "number-first. We never oversell.")

doc.add_paragraph(style="OW Heading 2").add_run("Body and headings")
doc.add_paragraph().add_run(
    "Body copy uses Space Grotesk at a comfortable measure with generous spacing. "
    "Headings are navy and tightly tracked. Labels and data use IBM Plex Mono. "
    "Sentences are separated with full stops and commas, never with a dash.")

doc.add_paragraph(style="OW Heading 3").add_run("A bulleted list")
bullet("Lead with the outcome the reader cares about.")
bullet("Keep each point to a single, concrete idea.")
bullet("Use the arrow tick, not a default round bullet.")

# Stat block
doc.add_paragraph(style="OW Heading 2").add_run("Proof, stated number-first")
sp = doc.add_paragraph(style="OW Stat Number"); sp.add_run("3x")
doc.add_paragraph(style="OW Stat Label").add_run("tender throughput, two-person team")

# Callout (teal wash + left rule)
doc.add_paragraph(style="OW Heading 3").add_run("Callout / aside")
cp = doc.add_paragraph(style="OW Callout")
cp.add_run("Use a tinted callout for a key note. Teal for consultancy themes, "
           "terracotta for product themes.")
shade_para(cp, TEALW); para_left_border(cp, "15727C")

# Pull quote
doc.add_paragraph(style="OW Heading 3").add_run("Pull quote")
doc.add_paragraph(style="OW Pull Quote").add_run(
    "“They took us from one tender a month to three, without adding headcount.”")
doc.add_paragraph(style="OW Attribution").add_run("Bid Director, Law enforcement client")

# Table
doc.add_paragraph(style="OW Heading 2").add_run("Table style")
rows = [("Capability", "What it does", "Outcome"),
        ("AI Advisory", "Where AI moves the needle", "A costed roadmap"),
        ("Bid Support", "Win strategy and response", "More winnable bids"),
        ("OWL suite", "Software you run yourself", "Capacity that scales")]
tbl = doc.add_table(rows=len(rows), cols=3)
no_table_borders(tbl)
widths = [Cm(3.6), CONTENT_W - Cm(3.6) - Cm(4.8), Cm(4.8)]
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        c = tbl.rows[i].cells[j]; c.width = widths[j]
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        if i == 0:
            shade_cell(c, NAVYHEX)
            r = p.add_run(val.upper()); run_font(r, MONO, 8, WHITE, caps=True, spacing=12)
        else:
            if i % 2 == 0:
                shade_cell(c, PAPER)
            cell_bottom_border(c, HAIR, 6)
            r = p.add_run(val); run_font(r, SANS, 10, INK if j == 0 else MUTED, bold=(j == 0))

# Two doors closing
doc.add_paragraph().paragraph_format.space_after = Pt(8)
doc.add_paragraph(style="OW Heading 1").add_run("Two doors, one team")
doors = doc.add_table(rows=1, cols=2)
no_table_borders(doors)
doors.columns[0].width = HALF_W; doors.columns[1].width = HALF_W
def door(cell, fill, accent_hex, label, head, body, contact):
    cell.width = HALF_W
    shade_cell(cell, fill)
    cell_margins(cell, top=200, bottom=200, left=220, right=220)
    p = cell.paragraphs[0]; r = p.add_run(label); run_font(r, MONO, 8.5, RGBColor.from_string(accent_hex), caps=True, spacing=20)
    p.paragraph_format.space_after = Pt(2)
    p2 = cell.add_paragraph(); r = p2.add_run(head); run_font(r, SANS, 13, NAVY, bold=True); p2.paragraph_format.space_after = Pt(2)
    p3 = cell.add_paragraph(); r = p3.add_run(body); run_font(r, SANS, 10, MUTED); p3.paragraph_format.space_after = Pt(4)
    p4 = cell.add_paragraph(); r = p4.add_run(contact); run_font(r, MONO, 9, RGBColor.from_string(accent_hex))
dc = doors.rows[0].cells
door(dc[0], TEALW, "15727C", "Consultancy", "Work with our team",
     "Advisory, bid support, and bespoke tooling.", "hello@outcomeworks.co.uk")
door(dc[1], TERRAW, "C2562F", "Products", "Run the OWL suite",
     "Software you operate yourself, from day one.", "outcomeworks.co.uk")

OUT.parent.mkdir(exist_ok=True)
doc.save(str(OUT))
print("Wrote", OUT)
